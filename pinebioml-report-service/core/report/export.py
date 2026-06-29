import os
import io
import base64
from typing import Dict, Any, List
import sys

# Suppress stderr writes during weasyprint import to avoid noisy console prints on systems without GTK
_stderr = sys.stderr
try:
    with open(os.devnull, 'w') as f:
        sys.stderr = f
        from weasyprint import HTML
except Exception:
    HTML = None
finally:
    sys.stderr = _stderr

from docx import Document
from docx.shared import Inches, Pt
from supabase import create_client, Client
from dotenv import load_dotenv

load_dotenv()

def export_pdf(html_content: str, output_path: str) -> str:
    """
    Converts rendered HTML content to a PDF using WeasyPrint.
    """
    # Weasyprint needs a base URL if there are relative paths, but our template embeds CSS
    print(f"Generating PDF report: {output_path}")
    HTML(string=html_content).write_pdf(output_path)
    return output_path

def export_docx(parsed_data: Dict[str, Any], llm_analysis: Dict[str, Any], images_dict: Dict[str, str], output_path: str) -> str:
    """
    Builds a structured DOCX report using python-docx.
    """
    print(f"Generating DOCX report: {output_path}")
    doc = Document()
    
    # Title
    dataset_name = parsed_data.get("dataset_name", "Dataset")
    title = doc.add_heading(f'PineBioML Analysis Report - {dataset_name}', 0)
    title.alignment = 1 # Center
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(llm_analysis.get("executive_summary", "No executive summary provided."))
    
    # Key Findings
    key_findings = llm_analysis.get("key_findings", [])
    if key_findings and isinstance(key_findings, list):
        doc.add_heading('Key Findings', level=2)
        for finding in key_findings:
            doc.add_paragraph(finding, style='List Bullet')
            
    # Best Model Analysis
    doc.add_heading('Best Model Analysis', level=1)
    doc.add_paragraph(llm_analysis.get("best_model_analysis", "No analysis provided."))
    
    # Statistical Interpretation & Overfitting
    doc.add_heading('Statistical Interpretation', level=2)
    doc.add_paragraph(llm_analysis.get("statistical_interpretation", ""))
    
    overfit = llm_analysis.get("overfitting_warning")
    if overfit and overfit.lower() != "none":
        doc.add_heading('Overfitting Warning', level=3)
        p = doc.add_paragraph()
        p.add_run(overfit).bold = True
        
    # Feature Analysis
    doc.add_heading('Feature Importance (SHAP)', level=1)
    doc.add_paragraph(llm_analysis.get("feature_analysis", ""))
    
    # Images / Visualization Gallery
    doc.add_heading('Visualization Gallery', level=1)
    for plot_type, b64_str in images_dict.items():
        doc.add_heading(plot_type.replace("_", " ").title(), level=2)
        try:
            image_bytes = base64.b64decode(b64_str)
            image_stream = io.BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(5.0))
        except Exception as e:
            doc.add_paragraph(f"[Error loading image for {plot_type}: {e}]")
            
    # Limitations
    doc.add_heading('Limitations', level=1)
    doc.add_paragraph(llm_analysis.get("limitations", ""))
    
    # Layperson Summary
    doc.add_heading('Layperson Summary (zh-TW)', level=1)
    doc.add_paragraph(llm_analysis.get("layperson_summary_zh_tw", ""))
    
    doc.save(output_path)
    return output_path

def init_supabase() -> Client | None:
    url = os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_KEY")
    if not url or not key:
        return None
    return create_client(url, key)

def upload_to_storage(file_path: str, bucket_name: str, object_name: str) -> str:
    """
    Uploads a file to Supabase storage and returns the public URL.
    Falls back to returning local path if Supabase is not configured.
    """
    supabase = init_supabase()
    if not supabase:
        print(f"Warning: Supabase credentials not found. Treating as local storage: {file_path}")
        return f"local://{file_path}"
        
    print(f"Uploading {object_name} to Supabase bucket '{bucket_name}'...")
    with open(file_path, 'rb') as f:
        file_bytes = f.read()
        
    try:
        supabase.storage.from_(bucket_name).upload(
            file=file_bytes,
            path=object_name,
            file_options={"content-type": "application/octet-stream"}
        )
        # Get public URL
        res = supabase.storage.from_(bucket_name).get_public_url(object_name)
        return res
    except Exception as e:
        print(f"Error uploading to Supabase: {e}")
        return f"local://{file_path}"

def register_report_in_db(metadata: Dict[str, Any], urls: Dict[str, str]) -> Dict[str, Any]:
    """
    Registers the generated report URLs and metadata in the Supabase database.
    """
    supabase = init_supabase()
    
    record = {
        "dataset_name": metadata.get("dataset_name"),
        "task_type": metadata.get("task_type"),
        "best_model": metadata.get("best_model"),
        "html_url": urls.get("html"),
        "pdf_url": urls.get("pdf"),
        "docx_url": urls.get("docx")
    }
    
    if not supabase:
        print(f"Warning: Supabase credentials not found. DB Registration mocked: {record}")
        return record
        
    print("Registering report links in Supabase database...")
    try:
        data = supabase.table("reports").insert(record).execute()
        return data.data[0] if data.data else record
    except Exception as e:
        print(f"Error registering report in DB: {e}")
        return record
